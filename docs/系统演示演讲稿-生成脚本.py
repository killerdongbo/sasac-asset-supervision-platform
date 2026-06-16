#!/usr/bin/env python
"""Generate a Word document demo script for the 国资监管平台 system presentation."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# Helper functions
def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = '黑体'
    p.space_after = Pt(4)

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.font.name = '宋体'
    p.space_after = Pt(20)

def add_h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = '黑体'
    p.space_before = Pt(20)
    p.space_after = Pt(10)

def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = '黑体'
    p.space_before = Pt(14)
    p.space_after = Pt(6)

def add_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.space_after = Pt(6)

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.name = '宋体'

def add_separator():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— — —')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(180, 180, 180)

# =====================================================================
# DOCUMENT CONTENT
# =====================================================================

add_title('湛江市国资监管平台')
add_subtitle('系统功能演示与汇报')
add_subtitle(f'演示日期：{datetime.now().strftime("%Y年%m月%d日")}')

add_separator()

# ---- 1. 演示概述 ----
add_h1('一、演示概述')

add_h2('1.1 演示目标')
add_body('本次演示旨在向领导全面展示湛江市国资监管平台（二期）的核心功能、数据架构和业务价值，重点突出国资监管"数字大脑"的穿透式监管能力，以及五大集团的组织架构重塑和资产全景可视化管理。')

add_h2('1.2 演示场景')
add_bullet('场景一：管理驾驶舱 — 国资监管数字大脑，一屏统览全市国资家底')
add_bullet('场景二：组织架构 — 五大集团及216家子公司股权穿透全景图')
add_bullet('场景三：资产管理 — 405项资产全生命周期数字化管理')
add_bullet('场景四：现场管理 — 巡检、盘点、维保全流程闭环')
add_bullet('场景五：审批与预警 — 智能化工作流审批+风险实时监控')
add_bullet('场景六：统计报表 — 19种国资监管标准报表一键导出')

add_h2('1.3 演示数据')
add_bullet('5大集团 + 216家各级子公司')
add_bullet('405项资产，总值约115.57亿元')
add_bullet('覆盖土地、房屋、设备、存货、无形资产、基础设施等10大资产类别')
add_bullet('114条折旧记录、6条巡检异常、4条预警记录')

add_separator()

# ---- 2. 系统架构 ----
add_h1('二、系统架构简述')

add_h2('2.1 技术架构')
add_bullet('前端：Vue 3 + TypeScript + Element Plus + ECharts 数据可视化')
add_bullet('后端：Spring Boot 3.2 + MyBatis-Plus + PostgreSQL + Redis')
add_bullet('文件存储：MinIO 对象存储（S3兼容）')
add_bullet('部署方式：Docker Compose 容器化一键部署')
add_bullet('安全体系：Spring Security + JWT 无状态认证 + 多租户数据隔离')

add_h2('2.2 业务架构')
add_body('系统以"国资委 → 五大集团 → 各级子公司"三级穿透式监管为核心数据模型，涵盖资产管理、现场管理、财务折旧、审批工作流、预警监控、审计追溯、统计报表七大业务模块，形成完整的国资监管闭环。')

add_h2('2.3 数据流转')
add_body('资产数据通过Excel导入或手工录入进入系统，经审批工作流确认后纳入台账管理。巡检和盘点模块定期扫描资产状态，发现异常自动触发维修工单或预警通知。折旧模块按月自动计提，财务数据直接对接监管报表。所有数据变更均通过AOP切面记录审计日志，确保全链路可追溯。')

add_separator()

# ---- 3. 演示流程 ----
add_h1('三、演示流程（约30分钟）')

# Scene 1
add_h2('🎬 场景一：管理驾驶舱（3分钟）')
add_body('【开场】各位领导好，欢迎视察湛江市国资监管平台。首先我们进入管理驾驶舱，这是我们系统的"数字大脑"—— 一屏统览全市国资家底。')
add_body('【演示操作】登录系统（admin / admin123），自动进入Dashboard页面。')
add_bullet('顶部四个核心指标卡片：资产总数405项、资产原值3.45亿元、本月新增31项、资产净值115.57亿元')
add_bullet('左侧饼图展示13个资产类别的分布情况，可以直观看到城市发展集团在房屋建筑和土地类的资产占比较大')
add_bullet('右侧展示使用状态分布：在册35项、闲置6项、出租2项、处置2项，闲置资产一目了然')
add_bullet('底部趋势图展示近12月资产新增趋势，反映集团资产规模的增长动态')
add_body('【话术要点】"这个页面让领导不用翻阅任何纸质报表，打开系统就能实时掌握全市国资的全景数据。每一个数字背后都有详细的数据支撑，我们可以随时下钻到具体的资产明细。"')

add_separator()

add_h2('🎬 场景二：组织架构与集团穿透（4分钟）')
add_body('点击左侧菜单"系统管理 → 组织架构"，展示五大集团的完整组织树。')
add_bullet('湛江市国有资产监督管理委员会（根节点）')
add_bullet('→ 湛江市城市发展集团有限公司（26家子公司）：城发建筑、高铁新城、西城新区、海上风电、科技创新园等')
add_bullet('→ 湛江市海洋与农业投资集团有限公司（16家子公司）：蓝色海洋产业、金海粮油、食品冷链等')
add_bullet('→ 湛江市旅游投资集团有限公司（17家子公司）：蓝月湾温泉、南极村、国际会展中心、军博园等')
add_bullet('→ 湛江市公共服务集团有限公司（8家子公司）：公交集团、公服巴士、汽车维修中心等')
add_bullet('→ 湛江市资产运营集团有限公司（149家子公司）：吉民药业、八建集团、湛江玻璃厂等')
add_body('点击"集团穿透统计"菜单，可以按集团维度查看各自的资产汇总和子公司分布。')
add_body('【话术要点】"我们根据爱企查的股权穿透数据，完整构建了五大集团及其全部216家各级子公司的组织架构。任何一家子公司的资产变动，都能自动汇总到集团层面，实现真正的穿透式监管。"')

add_separator()

add_h2('🎬 场景三：资产管理（5分钟）')
add_body('点击"资产台账 → 资产清单"，进入资产管理核心页面。')
add_bullet('支持按关键词、资产类别、使用状态多维度筛选')
add_bullet('分页展示405项资产，显示资产编号、名称、类别、原值、净值、累计折旧、使用状态、所在地')
add_bullet('支持新增资产、Excel批量导入、查看详情、编辑、删除')
add_bullet('点击任意资产可查看完整档案：基本信息、财务折旧、巡检记录、操作日志')
add_body('【演示操作】')
add_bullet('筛选"土地类"资产，展示3宗地块的详细信息，包括坐落位置、面积、权属证号')
add_bullet('筛选"闲置"状态的资产，展示6项闲置资产，其中霞山商务区B地块已闲置3年，系统已生成巡检异常')
add_bullet('点击"新增资产"，演示快速录入手持设备登记')
add_body('【话术要点】"系统实现了资产的\'一物一码\'全生命周期管理。从采购验收、入库登记、日常使用、巡检维保到最终报废处置，每个环节都有据可查、有迹可循。"')

add_separator()

add_h2('🎬 场景四：现场管理 — 巡检与盘点（4分钟）')
add_body('点击"现场管理 → 巡检管理"，展示巡检任务列表。')
add_bullet('巡检任务支持按组织和资产类型筛选范围，自动生成巡检清单')
add_bullet('巡检执行页面支持扫码确认、拍照上传、异常登记')
add_bullet('巡检异常自动归类：闲置浪费、报废需求、库存积压等')
add_bullet('异常记录可一键转维保工单，形成"发现→处理→验收"闭环')
add_body('【演示操作】点击"盘点管理 → 盘点任务"，展示盘点差异自动比对功能。系统在盘点过程中实时比对账实数据，差异项自动高亮标记。')
add_body('【话术要点】"传统的巡检盘点靠纸质表单和Excel汇总，效率低、易出错。我们的系统支持移动端扫码巡检，异常自动上报、自动流转，大幅提升了现场管理效率和数据准确性。"')

add_separator()

add_h2('🎬 场景五：审批工作流与预警监控（4分钟）')
add_body('点击"审批工作流 → 待我审批"，展示待审批事项列表。')
add_bullet('支持多级审批链配置（直属上级 → 部门负责人 → 集团分管领导 → 国资委终审）')
add_bullet('审批流程可视化，可查看当前节点和审批历史')
add_bullet('支持超时自动升级（如实例400节点1超时升级到ORG_MANAGER）')
add_body('点击"预警与审计 → 预警中心"，展示系统自动检测的风险项。')
add_bullet('闲置资产预警：霞山商务区B地块闲置超3年')
add_bullet('报废需求预警：旧办公楼楼龄超20年结构安全隐患')
add_bullet('呆滞库存预警：旧型号轴承库存积压超2年')
add_body('点击"预警与审计 → 操作日志"，所有数据变更均有AOP自动记录的审计日志，支持按操作人、操作类型、时间范围查询。')
add_body('【话术要点】"审批工作流实现了\'让数据多跑路，让人少跑腿\'。预警系统则像24小时不间断的电子巡检员，主动发现问题、主动推送预警，把风险消灭在萌芽状态。"')

add_separator()

add_h2('🎬 场景六：统计报表与数据导出（5分钟）')
add_body('点击"统计报表 → 📥 报表导出"，进入报表导出中心。')
add_body('【演示操作 — Tab 1：报表导出】')
add_bullet('展示19种国资监管标准报表：资产负债表、资产底数清单、问题资产及整治清单、盘活利用清单，以及15类资产清查明细表和勾稽关系稽核表')
add_bullet('选择"资产负债表"→ 选择周期"2026-06" → 点击"导出Excel"')
add_bullet('任务状态变为"处理中"→ "已完成"后，点击"下载"即可获取标准Excel报表')
add_body('【演示操作 — Tab 2：模板下载】')
add_bullet('选择任意报表类型 → 点击"下载模板"，获取含表头和填写说明的标准模板')
add_bullet('说明：基层单位可以下载模板→线下填报→上传导入，实现数据采集标准化')
add_body('【演示操作 — Tab 3：报表导入】')
add_bullet('上传按模板填好的Excel → 系统自动校验列匹配和数据格式 → 预览确认 → 一键入库')
add_bullet('校验失败会精确提示第几行、哪个字段、什么原因')
add_body('【话术要点】"这19种报表覆盖了国资委要求的全部上报类型。以前需要手工汇总、反复核对，现在一键导出，而且每张报表之间通过勾稽关系表自动校验数据一致性，确保上报数据的准确性和合规性。"')

add_separator()

add_h2('🎬 补充场景：财务折旧（2分钟）')
add_body('点击"财务折旧 → 折旧管理"，展示按月自动计提的折旧明细。')
add_bullet('系统按资产类别和折旧方法（直线法、双倍余额递减法等）自动计算每月折旧')
add_bullet('支持按月份、资产筛选，展示折旧前值、折旧额、折旧后值')
add_bullet('当前系统记录114条折旧明细，5月折旧总额约95.65万元')
add_body('选择"统计报表 → 折旧分析"，展示按组织和月份的折旧趋势图。')

add_separator()

# ---- 4. 技术亮点 ----
add_h1('四、系统技术亮点')

add_h2('4.1 穿透式监管')
add_body('基于MyBatis-Plus的递归SQL查询，实现"国资委→集团→子公司"三级组织树下的资产数据自动汇总。任何子公司的资产变动实时反映到集团和国资委层面的统计报表。')

add_h2('4.2 异步任务引擎')
add_body('报表导出采用Spring @Async异步任务+MinIO对象存储架构。大数据量报表（如资产底数清单导出全部405条资产记录）不会阻塞用户操作，任务后台执行完成后自动通知下载。')

add_h2('4.3 AOP审计追踪')
add_body('所有数据变更（新增、修改、删除）均通过Spring AOP切面自动记录操作人、操作时间、操作内容、IP地址，确保全链路可审计、可追溯，满足国资监管的合规要求。')

add_h2('4.4 多租户数据隔离')
add_body('系统支持多租户架构，每个集团或组织通过X-Tenant-Id请求头实现数据隔离。同一套系统可同时服务于多个独立组织，数据互不干扰。')

add_h2('4.5 容器化部署')
add_body('系统采用Docker Compose容器化部署，PostgreSQL + Redis + MinIO + Spring Boot + Nginx五大容器一键启动。支持在任何Linux/Windows服务器上快速部署，运维成本极低。')

add_separator()

# ---- 5. 数据总览 ----
add_h1('五、系统核心数据一览')

add_body('以下数据供演示过程中随时引用：')

# Table
table = doc.add_table(rows=12, cols=3, style='Light Grid Accent 1')
table.autofit = True

headers = ['指标', '数值', '说明']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(11)

data = [
    ['五大集团', '湛江市城市发展集团、海洋与农业投资集团、旅游投资集团、公共服务集团、资产运营集团', ''],
    ['各级子公司', '216家', '根据爱企查股权穿透数据'],
    ['资产总数', '405项', '覆盖10大资产类别'],
    ['资产总值', '约115.57亿元（原值）', '城市发展集团占比最高（66.67亿）'],
    ['城市发展集团', '97项资产，¥66.67亿', '含5项地标资产'],
    ['海洋与农业投资集团', '31项资产，¥5.42亿', ''],
    ['旅游投资集团', '30项资产，¥4.11亿', ''],
    ['公共服务集团', '16项资产，¥1.94亿', ''],
    ['资产运营集团', '229项资产，¥37.12亿', '含149家子公司'],
    ['折旧记录', '114条', '覆盖2025年8月至2026年6月'],
    ['监管报表', '19种标准报表', '覆盖国资委全部上报类型'],
]

for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        cell = table.rows[i+1].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = '宋体'

add_separator()

# ---- 6. 演示准备 ----
add_h1('六、演示环境准备清单')

add_bullet('服务器：确保Docker环境正常，所有容器运行中（docker ps 确认5个容器状态）')
add_bullet('网络：演示电脑可访问 http://localhost （前端）和 http://localhost:8082（后端API）')
add_bullet('浏览器：推荐Chrome/Edge最新版，清除缓存后登录')
add_bullet('登录账号：admin / admin123')
add_bullet('演示数据：确认资产列表显示405项，组织架构显示222个节点')
add_bullet('备用方案：如网络异常，可使用localhost本地演示')
add_bullet('导出测试：提前创建1-2个报表导出任务，演示时可直接下载已完成文件')
add_bullet('MinIO控制台：如需演示文件存储，可访问 http://localhost:9001（minioadmin/minioadmin）')

add_separator()

# ---- 7. FAQ ----
add_h1('七、常见问题预案')

add_h2('Q1：系统数据安全性如何保障？')
add_body('系统采用Spring Security + JWT无状态认证，所有API接口均需携带有效Token访问。数据层面通过多租户隔离和AOP审计日志确保数据安全。生产环境建议配置HTTPS和数据库加密。')

add_h2('Q2：能否对接现有财务系统？')
add_body('系统使用标准REST API接口，可通过定制开发对接SAP、用友、金蝶等主流财务系统。折旧模块已预留财务数据导入接口。')

add_h2('Q3：系统支持多少并发用户？')
add_body('基于Spring Boot的异步架构和PostgreSQL的连接池管理，单机部署可支持200+并发用户。报表导出采用异步任务队列，不会因单次大数据量导出影响其他用户体验。')

add_h2('Q4：数据如何备份？')
add_body('PostgreSQL和MinIO数据均通过Docker Volume持久化存储。生产环境建议配合pg_dump定时备份和MinIO镜像复制策略。')

add_h2('Q5：系统后续迭代计划？')
add_body('计划支持移动端APP巡检、AI智能估值、多维度数据大屏、国资委直报系统对接等功能。欢迎领导提出宝贵意见和建议。')

add_separator()

# ---- End ----
add_h1('八、结束语')
add_body('感谢各位领导在百忙之中抽出时间听取汇报。湛江市国资监管平台的建设，是落实国资委关于推进国资监管数字化、智能化转型的重要举措。我们将持续优化系统功能，为湛江市国资监管工作提供更加有力的技术支撑。')
add_body('请各位领导批评指正！')

# ---- Save ----
import os
output_filename = '湛江市国资监管平台-系统演示演讲稿.docx'
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), output_filename)
doc.save(output_path)
print(f'Document saved to: {output_path}')
print(f'File size: {os.path.getsize(output_path)} bytes')

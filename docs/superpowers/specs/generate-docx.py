#!/usr/bin/env python3
"""生成包含彩色架构图的功能设计Word文档"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

BASE = os.path.dirname(os.path.abspath(__file__))
DIAGRAMS = os.path.join(BASE, 'diagrams')
SPEC_FILE = os.path.join(BASE, '2026-06-12-湛江国资监管平台二期-完整功能设计-v2.md')
OUTPUT = os.path.join(BASE, '2026-06-12-湛江国资监管平台二期-完整功能设计-v2.docx')

doc = Document()

# === Styles ===
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)

for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

def add_title(text):
    h = doc.add_heading(text, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

def add_h1(text): return doc.add_heading(text, level=1)
def add_h2(text): return doc.add_heading(text, level=2)
def add_h3(text): return doc.add_heading(text, level=3)

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Light Grid Accent 1')
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()
    return table

def add_image(name, width=6.0):
    path = os.path.join(DIAGRAMS, name)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(path, width=Inches(width))
        doc.add_paragraph()

# ===== BUILD DOCUMENT =====

add_title('湛江国资国企一体化数字平台（二期）\n完整功能设计文档 v2')
add_para('2026-06-12 | 基于功能清单（二期）.xlsx 完善', bold=False)
doc.add_paragraph()

# -- Section 1 --
add_h1('一、修订说明')
add_para('本文档在 v1 设计文档基础上，依据《湛江市国资国企一体化数字平台功能清单（二期）.xlsx》的60项功能清单进行了全面补充和细化。')
add_para('主要变更：')
changes = ['模块从 5个业务服务 扩展为 16个功能模块、60+功能点',
           '新增：巡检管理、盘点管理、采购管理、维保维修管理、预警中心',
           '补充：资产基础资料（位置/供应商/维保商）、资产标签、操作日志',
           '审批工作流从笼统描述升级为可配置审批引擎',
           '交付周期从 7周 调整为 14周两阶段交付']
for c in changes:
    doc.add_paragraph(c, style='List Bullet')

# -- Section 2 --
add_h1('二、关键设计决策')
add_table(
    ['#', '决策项', '选择', '说明'],
    [['1', '覆盖范围', '功能清单16模块全部纳入', '企业内资产管理 + 国资监管上报全覆盖'],
     ['2', '交付周期', '14周两阶段（8-9月）', 'Phase 1(7月底)：台账+流转+采购+折旧+驾驶舱；Phase 2(8-9月)：巡检+盘点+维保+审批+预警+监管上报'],
     ['3', '系统架构', '模块化单体 + 三域划分', '企业资产管理域 / 国资监管域 / 平台基础域'],
     ['4', '企业层vs监管层', '平行独立、数据共享', '企业域数据作为监管上报数据源，通过接口松耦合'],
     ['5', '审批工作流', '可配置审批引擎', '按业务类型定义审批链，支持条件分支'],
     ['6', '资产管理深度', '全生命周期', '从采购申请到报废退出的完整闭环'],
     ['7', '数据采集策略', '混合模式', 'Excel批量导入 + 财务系统对接 + 人工录入'],
     ['8', '部署模式', 'SaaS多租户', '统一部署，tenant_id数据隔离'],
     ['9', '技术栈', 'Java/SpringBoot全栈', 'SpringBoot 3 + Vue3 + Element Plus + PostgreSQL']]
)

# -- Section 3: Diagrams --
add_h1('三、功能模块全景图')

add_h2('3.1 系统三域架构')
add_para('系统划分为三个业务域，域内模块高内聚，域间通过明确的 Service 接口通信：')
add_image('01-three-domains.png', width=6.2)

add_h2('3.2 域间关系说明')
add_para('企业域 → 监管域：企业资产台账通过数据接口提供给监管上报服务，生成19类报表快照。企业域不感知监管上报的存在，松耦合。')
add_para('审批工作流跨域共用：审批引擎作为独立模块，服务于调拨审批（企业域）、盘点差异审批（企业域）、上报审核（监管域）。')
add_para('平台域横切能力：多租户隔离、RBAC权限、操作日志、数据导入导出由平台域统一提供，其他域自动继承。')

add_h2('3.3 完整功能清单')
add_table(
    ['序号', '模块', '功能', '域'],
    [['1', '系统入口与安全', '用户登录、个人身份识别、访问控制', '平台基础'],
     ['2', '组织权限管理', '用户管理、角色管理、权限管理、部门管理', '平台基础'],
     ['3', '数据导入导出', '资产批量导入、清单与报告导出', '平台基础'],
     ['4', '资产基础资料', '资产分类、存放位置、供应商档案、维保商档案', '企业资产'],
     ['5', '资产台账管理', '资产登记、自动编号、清单查询、信息维护、详情卡片、标签管理、变更记录', '企业资产'],
     ['6', '资产流转管理', '入库管理、资产领用、资产调拨、资产报废、流转单据查询', '企业资产'],
     ['7', '采购管理', '采购申请、采购验收、采购转资产', '企业资产'],
     ['8', '巡检管理', '巡检任务制定、任务跟踪、现场登记、我的任务、异常处理', '企业资产'],
     ['9', '盘点管理', '盘点任务制定、任务跟踪、现场登记、我的任务、差异识别、差异报告', '企业资产'],
     ['10', '维保维修管理', '维保申请、维修登记、巡检异常转维修', '企业资产'],
     ['11', '财务折旧管理', '折旧规则设置、折旧计提、折旧确认', '企业资产'],
     ['12', '工作台', '管理驾驶舱、资产结构分析、近期业务动态', '国资监管'],
     ['13', '统计报表', '资产统计报表、巡检与盘点统计、折旧统计', '国资监管'],
     ['14', '监管上报', '19类报表采集、审核链（企业→集团→国资委）、版本快照', '国资监管'],
     ['15', '审批工作流', '待我审批、我发起的申请、我已审批', '国资监管'],
     ['16', '预警中心', '维保到期预警、巡检超期预警、借用超期预警、闲置资产预警', '国资监管'],
     ['17', '审计追溯', '操作日志、全生命周期追踪', '国资监管'],
     ['18', '财务对接', '财务系统接口适配、数据拉取、清洗校验、合并抵销校验', '国资监管']]
)

# -- Section 4: Data Model --
add_h1('四、核心数据模型')

add_h2('4.1 核心实体关系图')
add_image('02-er-diagram.png', width=6.2)

add_h2('4.2 企业资产管理域 — 核心实体')
add_table(
    ['模块', '核心实体', '关键字段', '关系'],
    [['资产基础资料', 'AssetCategory/Location/Supplier/MaintenanceProvider', 'code,name,parent_id/name,address/name,contact', '独立字典表'],
     ['资产台账', 'Asset/AssetLabel/AssetChangeLog', 'name,asset_code,category,org_id,value,label_code', '1:N关系'],
     ['资产流转', 'StockIn/Assignment/Transfer/Disposal', 'asset_id,from_org,to_org,operator,status', '流转单→Asset(N:1)'],
     ['采购管理', 'PurchaseRequest/PurchaseAcceptance', 'asset_name,quantity,budget,supplier_id', 'Req→Accept(1:1)→Asset'],
     ['巡检管理', 'InspectionTask/InspectionRecord/InspectionAnomaly', 'task_name,assignee,asset_scope(JSON),anomaly_type', 'Task→Records(1:N)'],
     ['盘点管理', 'InventoryTask/InventoryRecord/InventoryDiff', 'task_name,scope,book_qty,actual_qty,diff_type', 'Task→Records(1:N)'],
     ['维保维修', 'MaintenanceRequest/MaintenanceRecord', 'asset_id,fault_description,provider_id,cost', 'Req→Record(1:1)'],
     ['折旧管理', 'DepreciationRule/DepreciationRecord', 'asset_id,period,amount,accumulated,confirmed', 'Record→Asset(N:1)']]
)

add_h2('4.3 国资监管域 — 核心实体')
add_table(
    ['模块', '核心实体', '关键字段'],
    [['审批工作流', 'ApprovalDef/ApprovalInstance/ApprovalNode', 'def_name,biz_type,trigger_condition(JSON)/instance_status,current_node'],
     ['监管上报', 'ReportTemplate/ReportSnapshot/ReportSubmission', 'report_type,period,org_id,snapshot_data(JSONB),submit_status'],
     ['预警中心', 'AlertRule/AlertRecord', 'alert_type,rule_config(JSON),enabled/triggered_at,alert_data(JSONB)'],
     ['审计追溯', 'OperationLog/AssetLifecycleView', 'operator_id,action,target_type,before_data(JSONB),after_data(JSONB)'],
     ['财务对接', 'FinancialData/FinancialAdapterConfig', 'org_id,adapter_type,connection_config(JSON),last_sync_at']]
)

# -- Section 5: Approval --
add_h1('五、审批工作流设计')
add_h2('5.1 审批状态机')
add_image('04-process-flows.png', width=6.2)

add_h2('5.2 支持的审批业务类型')
add_table(
    ['业务类型', '默认审批链', '条件分支'],
    [['资产调拨', '调入部门确认', '—'],
     ['资产报废', '部门负责人 → 财务审核', '原值>10万：追加分管领导'],
     ['盘点差异', '差异确认 → 财务调整', '—'],
     ['采购申请', '部门负责人 → 管理层', '金额>5万：追加财务审核'],
     ['维保维修', '申请 → 部门确认', '—'],
     ['监管上报', '企业提交 → 集团审核 → 国资委接收', '—']]
)

# -- Section 6: Processes --
add_h1('六、核心业务流程')
add_para('以下流程图展示了巡检管理、盘点管理、采购管理、维保维修管理、监管上报、预警中心、审批工作流、资产全生命周期等八大核心业务流程。')
add_para('（详细流程请参见上方"核心业务流程图"）')

# -- Section 7 & 8 --
add_h1('七、多租户与权限模型')
add_table(
    ['维度', '方案'],
    [['数据库级', '共享数据库 + tenant_id 字段隔离'],
     ['应用级', 'TenantContext（ThreadLocal）自动注入，MyBatis-Plus 拦截器统一过滤'],
     ['组织级', '树形层级权限：上级可查看下级数据，同级隔离'],
     ['角色级', 'RBAC 四角色（国资委监管员/集团审核员/企业操作员/系统管理员）+ 可扩展自定义角色']]
)

add_h1('八、技术选型')
add_table(
    ['层次', '选型', '理由'],
    [['前端框架', 'Vue 3 + TypeScript + Element Plus', '政务项目主流选型，组件库成熟'],
     ['可视化/驾驶舱', 'ECharts 5', '国产可视化库，大屏支持成熟'],
     ['后端框架', 'Spring Boot 3.x + MyBatis-Plus', '政务信创标配，生态完善'],
     ['数据库', 'PostgreSQL 15 + PostGIS', '开源、信创兼容，JSONB 支持灵活配置'],
     ['缓存', 'Redis 7', 'Session共享、租户配置缓存、预警结果缓存'],
     ['文件存储', 'MinIO', '资产照片、巡检拍照、附件、报表导出'],
     ['报表引擎', 'JasperReports + Apache POI', '复杂报表模板 + Excel导入导出'],
     ['定时任务', 'Spring @Scheduled', '折旧计提、预警扫描（MVP够用）'],
     ['容器化', 'Docker + Docker Compose', 'MVP够用，后续升级K8s'],
     ['信创适配', '达梦/人大金仓 + 麒麟OS + 东方通', '基于PG协议兼容，迁移成本低']]
)

# -- Section 9: Deployment --
add_h1('九、部署架构与代码组织')
add_image('05-code-deploy.png', width=6.2)

# -- Section 10: Timeline --
add_h1('十、分期交付计划（14周）')
add_image('03-delivery-timeline.png', width=6.2)

add_h2('10.1 Phase 1 验收标准')
add_para('企业可通过系统完成采购申请→资产入账→日常使用→折旧计提→报废退出的完整闭环，驾驶舱展示企业资产全貌。')

add_h2('10.2 Phase 2 验收标准')
add_para('企业可完成巡检盘点现场管理，数据可通过审批上报至国资委，系统具备完整的审计追溯能力。')

# -- Section 11: API --
add_h1('十一、API 设计原则')
add_para('统一响应格式：')
add_para('成功: { "success": true, "data": {...}, "meta": { "total": 100, "page": 1, "limit": 20 } }', bold=False)
add_para('错误: { "success": false, "error": "资产编码已存在: ZC-2026-0001" }', bold=False)
add_para('主要API端点：/api/assets、/api/inspection-tasks、/api/inventory-tasks、/api/reports、/api/approval-instances、/api/alerts 等。')
add_para('多租户：所有 /api/** 请求通过 X-Tenant-Id Header 注入租户上下文。')

# -- Section 12 & 13 --
add_h1('十二、错误处理策略')
add_table(
    ['错误类型', 'HTTP状态码', '处理方式'],
    [['参数校验失败', '400', 'GlobalExceptionHandler 统一拦截'],
     ['业务逻辑错误', '400', 'BusinessException，返回中文错误信息'],
     ['认证失败', '401', 'Spring Security 拦截'],
     ['权限不足', '403', '返回"无权限访问"'],
     ['资源不存在', '404', '返回"XX不存在"'],
     ['系统内部错误', '500', '记录详细日志，返回"系统内部错误"']]
)

add_h1('十三、测试策略')
add_para('目标覆盖率：≥ 80%')
add_table(
    ['测试类型', '范围', '工具'],
    [['单元测试', 'Service 层业务逻辑、工具类、审批引擎', 'JUnit 5 + Mockito'],
     ['集成测试', 'Controller API 端点、数据库操作、事务边界', 'Spring Boot Test + Testcontainers'],
     ['前端测试（Phase 2）', '关键页面交互、审批流程', 'Vitest + Vue Test Utils'],
     ['E2E测试（Phase 2）', '核心用户流程', 'Playwright']]
)

# -- Section 14 --
add_h1('十四、风险与缓解')
add_table(
    ['风险', '影响', '缓解措施'],
    [['14周范围大', '质量不达标', '分期交付，Phase 1聚焦核心闭环'],
     ['审批引擎复杂度高', 'Phase 1延期', 'Sprint 2先做固定链审批，Sprint 3再抽象'],
     ['财务对接难度高', '数据拿不到', '适配器层设计灵活，Excel模板兜底'],
     ['新模块测试不足', 'Bug多', 'TDD强制，每个Sprint检查覆盖率'],
     ['企业配合度不足', '数据进不来', 'Excel批量导入降低门槛，驻场培训'],
     ['信创环境兼容', '部署受阻', '技术栈信创兼容选型，MVP标准环境先跑']]
)

# -- Save --
doc.save(OUTPUT)
print(f'Word document saved to: {OUTPUT}')
